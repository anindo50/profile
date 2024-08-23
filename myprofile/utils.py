from pdfminer.high_level import extract_text

def pdf_to_text(pdf_path, txt_path):
    # Extract text from PDF
    text = extract_text(pdf_path)
    
    # Save the extracted text to a text file with UTF-8 encoding
    with open(txt_path, 'w', encoding='utf-8') as txt_file:
         txt_file.write(text)
    return txt_path


# def pdf_to_text(pdf_path, txt_path):
#     # Example implementation, replace with your actual conversion logic
#     with open(pdf_path, 'rb') as pdf_file:
#         extract_text(pdf_path)
    
#     # Save the extracted text to a text file with UTF-8 encoding
#     # with open(txt_path, 'w', encoding='utf-8') as txt_file:
#     #     txt_file = txt_file.write(text)
#     return txt_path

# Example usage
# pdf_to_text('C:/Users/User/Desktop/test_ai/book.pdf', 'C:/Users/User/Desktop/test_ai/book1.txt')

'''with pytesseract ocr'''


# import fitz  # PyMuPDF
# from PIL import Image
# import pytesseract
# from docx import Document
# import io

# # Path to the Tesseract executable
# # Update this path if Tesseract is not in your PATH environment variable
# pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

# def pdf_to_images(pdf_path):
#     pdf_document = fitz.open(pdf_path)
#     images = []
#     for page_num in range(len(pdf_document)):
#         page = pdf_document.load_page(page_num)
#         pix = page.get_pixmap()
#         img = Image.open(io.BytesIO(pix.tobytes()))
#         images.append(img)
#     return images

# def ocr_images(images):
#     text = ""
#     for img in images:
#         text += pytesseract.image_to_string(img).strip() + "\n\n"
#     return text

# def text_to_word(text, output_path):
#     doc = Document()
#     doc.add_paragraph(text)
#     doc.save(output_path)

# def pdf_to_word(pdf_path, output_path):
#     images = pdf_to_images(pdf_path)
#     text = ocr_images(images)
#     text_to_word(text, output_path)



import fitz  # PyMuPDF
from PIL import Image
import pytesseract
from docx import Document
import io
import cv2
import numpy as np

# Path to the Tesseract executable
# Update this path if Tesseract is not in your PATH environment variable
# pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

def pdf_to_images(pdf_path, dpi=300):
    """
    Convert PDF pages to images with specified DPI for better OCR accuracy.
    """
    pdf_document = fitz.open(pdf_path)
    images = []
    for page_num in range(len(pdf_document)):
        page = pdf_document.load_page(page_num)
        pix = page.get_pixmap(dpi=dpi)  # Set a higher DPI for better image resolution
        img = Image.open(io.BytesIO(pix.tobytes()))
        images.append(img)
    return images

def preprocess_image_with_opencv(img):
    """
    Preprocess the image using OpenCV to improve OCR accuracy.
    This involves converting to grayscale, applying thresholding, and noise reduction.
    """
    # Convert PIL Image to OpenCV format
    img_cv = np.array(img)
    img_cv = cv2.cvtColor(img_cv, cv2.COLOR_RGB2BGR)

    # Convert to grayscale
    gray = cv2.cvtColor(img_cv, cv2.COLOR_BGR2GRAY)

    # Apply Gaussian Blurring to reduce noise
    blurred = cv2.GaussianBlur(gray, (5, 5), 0)

    # Apply adaptive thresholding to make the text stand out
    thresh = cv2.adaptiveThreshold(blurred, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, 
                                   cv2.THRESH_BINARY, 11, 2)

    # Optionally, dilate or erode to improve character separation
    kernel = np.ones((1, 1), np.uint8)
    processed_img = cv2.dilate(thresh, kernel, iterations=1)

    # Convert back to PIL Image for Tesseract
    return Image.fromarray(processed_img)

def ocr_images(images, lang='eng'):
    """
    Perform OCR on the given images and extract text.
    Handles OCR errors and allows for language specification.
    """
    text = ""
    for img in images:
        try:
            # Preprocess the image using OpenCV
            processed_img = preprocess_image_with_opencv(img)
            
            # Extract text from the preprocessed image
            extracted_text = pytesseract.image_to_string(processed_img, lang=lang).strip()
            text += extracted_text + "\n\n"
        except pytesseract.TesseractError as e:
            # Handle any errors during OCR
            print(f"OCR error occurred: {e}")
            text += "[Error in OCR]\n\n"
    return text

def text_to_word(text, output_path):
    """
    Save the extracted text to a Word document.
    """
    doc = Document()
    doc.add_paragraph(text)
    doc.save(output_path)

def pdf_to_word(pdf_path, output_path, dpi=300, lang='eng'):
    """
    Full pipeline: Convert PDF to images, extract text via OCR, and save it to a Word document.
    Includes options for setting image DPI and OCR language.
    """
    images = pdf_to_images(pdf_path, dpi=dpi)
    text = ocr_images(images, lang=lang)
    # print(text)  # Optional: Print the text to the console for inspection
    text_to_word(text, output_path)



'''with google vision ocr'''
# from google.cloud import vision
# from google.oauth2 import service_account
# from PIL import Image
# import fitz  # PyMuPDF
# from docx import Document
# import io

# def pdf_to_images(pdf_path):
#     pdf_document = fitz.open(pdf_path)
#     images = []
#     for page_num in range(len(pdf_document)):
#         page = pdf_document.load_page(page_num)
#         pix = page.get_pixmap()
#         img = Image.open(io.BytesIO(pix.tobytes()))
#         images.append(img)
#     return images

# def ocr_images(images):
#     text = ""
#     client = vision.ImageAnnotatorClient()
#     for img in images:
#         buffer = io.BytesIO()
#         img.save(buffer, format='PNG')
#         content = buffer.getvalue()
#         image = vision.Image(content=content)
#         response = client.text_detection(image=image)
#         text += response.full_text_annotation.text + "\n\n"
#     return text

# def text_to_word(text, output_path):
#     doc = Document()
#     doc.add_paragraph(text)
#     doc.save(output_path)

# def pdf_to_word(pdf_path, output_path):
#     images = pdf_to_images(pdf_path)
#     text = ocr_images(images)
#     text_to_word(text, output_path)

'''with easyocr ocr
'''

# import easyocr
# from PIL import Image
# import fitz  # PyMuPDF
# from docx import Document
# import io

# def pdf_to_images(pdf_path):
#     pdf_document = fitz.open(pdf_path)
#     images = []
#     for page_num in range(len(pdf_document)):
#         page = pdf_document.load_page(page_num)
#         pix = page.get_pixmap()
#         img = Image.open(io.BytesIO(pix.tobytes()))
#         images.append(img)
#     return images

# def ocr_images(images):
#     text = ""
#     reader = easyocr.Reader(['en'])  # Specify the language(s) you need
#     for img in images:
#         text += " ".join(reader.readtext(img, detail=0))  # detail=0 returns only the text
#         text += "\n\n"
#     return text

# def text_to_word(text, output_path):
#     doc = Document()
#     doc.add_paragraph(text)
#     doc.save(output_path)

# def pdf_to_word(pdf_path, output_path):
#     images = pdf_to_images(pdf_path)
#     text = ocr_images(images)
#     text_to_word(text, output_path)


